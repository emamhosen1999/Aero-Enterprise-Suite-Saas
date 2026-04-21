#!/usr/bin/env python3
"""Simple DOCX <-> Markdown converter for skill workflows.

Usage:
  python docx_io.py extract --input input.docx --output input.md
  python docx_io.py render --input final.md --output output.docx --title "Final Report"
"""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable

from docx import Document


def _paragraph_to_markdown(style_name: str, text: str) -> str:
    cleaned = text.strip()
    if not cleaned:
        return ""

    if style_name.startswith("Heading"):
        # Common style names: Heading 1, Heading 2, ...
        parts = style_name.split()
        level = 1
        if len(parts) > 1 and parts[1].isdigit():
            level = max(1, min(6, int(parts[1])))
        return f"{'#' * level} {cleaned}"

    if "List Bullet" in style_name:
        return f"- {cleaned}"

    if "List Number" in style_name:
        return f"1. {cleaned}"

    return cleaned


def extract_docx_to_markdown(input_path: Path, output_path: Path) -> None:
    doc = Document(str(input_path))
    lines: list[str] = []

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else "Normal"
        md_line = _paragraph_to_markdown(style_name, paragraph.text)
        if md_line:
            lines.append(md_line)
            lines.append("")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def _is_numbered_item(line: str) -> bool:
    if ". " not in line:
        return False
    head, _ = line.split(". ", 1)
    return head.isdigit()


def _read_lines(path: Path) -> Iterable[str]:
    return path.read_text(encoding="utf-8").splitlines()


def render_markdown_to_docx(input_path: Path, output_path: Path, title: str | None) -> None:
    document = Document()

    if title:
        document.add_heading(title, level=0)

    for raw_line in _read_lines(input_path):
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("###### "):
            document.add_heading(stripped[7:].strip(), level=6)
            continue
        if stripped.startswith("##### "):
            document.add_heading(stripped[6:].strip(), level=5)
            continue
        if stripped.startswith("#### "):
            document.add_heading(stripped[5:].strip(), level=4)
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        if _is_numbered_item(stripped):
            _, item = stripped.split(". ", 1)
            document.add_paragraph(item.strip(), style="List Number")
            continue

        document.add_paragraph(stripped)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="DOCX/Markdown I/O helper")
    subparsers = parser.add_subparsers(dest="command", required=True)

    extract_parser = subparsers.add_parser("extract", help="Extract DOCX into Markdown")
    extract_parser.add_argument("--input", required=True, help="Input .docx path")
    extract_parser.add_argument("--output", required=True, help="Output .md path")

    render_parser = subparsers.add_parser("render", help="Render Markdown into DOCX")
    render_parser.add_argument("--input", required=True, help="Input .md path")
    render_parser.add_argument("--output", required=True, help="Output .docx path")
    render_parser.add_argument("--title", required=False, help="Optional document title")

    return parser.parse_args()


def main() -> None:
    args = parse_args()

    if args.command == "extract":
        extract_docx_to_markdown(Path(args.input), Path(args.output))
        return

    if args.command == "render":
        render_markdown_to_docx(Path(args.input), Path(args.output), args.title)
        return

    raise ValueError(f"Unsupported command: {args.command}")


if __name__ == "__main__":
    main()
